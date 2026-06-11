import os
import io
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

app = Flask(__name__)
CORS(app)

# Load .env file next to this script if present
_env_path = os.path.join(os.path.dirname(__file__), ".env")
if os.path.exists(_env_path):
    with open(_env_path) as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _v = _line.split("=", 1)
                os.environ.setdefault(_k.strip(), _v.strip())

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")


def expand_with_ai(section_name, short_notes, context=""):
    """Use Claude to expand short notes into professional CMO report language."""
    if not short_notes or not short_notes.strip():
        return ""

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    system = (
        "You are a Chief Marketing Officer writing a professional weekly marketing report. "
        "Expand the user's short bullet points or notes into polished, professional, concise prose "
        "suitable for a C-suite weekly report. Be specific, use marketing terminology, and sound confident. "
        "Return ONLY the expanded text with no preamble, no section headers, no extra commentary. "
        "Keep it to 2-4 sentences unless more is clearly needed. Do not use markdown formatting."
    )

    user_msg = f"Section: {section_name}\n"
    if context:
        user_msg += f"Additional context: {context}\n"
    user_msg += f"\nShort notes to expand:\n{short_notes}"

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=400,
        messages=[{"role": "user", "content": user_msg}],
        system=system,
    )
    return message.content[0].text.strip()


TNR = "Times New Roman"
BODY_PT = 12   # template default body size
TITLE_PT = 24  # sz=48 half-points
HEADER_PT = 18 # sz=36 half-points


def _apply_spacing(para):
    """Match template paragraph spacing: auto before/after, single line."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:beforeAutospacing"), "1")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:afterAutospacing"), "1")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    # replace existing spacing element if present
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def _tnr_run(para, text, bold=False, size_pt=BODY_PT, italic=False):
    run = para.add_run(text)
    run.font.name = TNR
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    # force east-asia / complex script font too
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn("w:ascii"), TNR)
        el.set(qn("w:eastAsia"), TNR)
        el.set(qn("w:hAnsi"), TNR)
        el.set(qn("w:cs"), TNR)
    return run


def _section_header(doc, number, title):
    """Bold 18pt Times New Roman section heading matching the template."""
    p = doc.add_paragraph()
    _apply_spacing(p)
    pPr = p._p.get_or_add_pPr()
    lvl = OxmlElement("w:outlineLvl")
    lvl.set(qn("w:val"), "1")
    pPr.append(lvl)
    _tnr_run(p, f"{number}. {title}", bold=True, size_pt=HEADER_PT)
    return p


def _bullet_item(doc, label, value):
    """Bullet list paragraph: bold label + plain value, matching template list style."""
    p = doc.add_paragraph(style="List Bullet")
    _apply_spacing(p)
    # override the list bullet font to Times New Roman
    if label:
        _tnr_run(p, f"{label}: ", bold=True, size_pt=BODY_PT)
    _tnr_run(p, value or "", bold=False, size_pt=BODY_PT)
    return p


def _plain_para(doc, text):
    """Plain body paragraph (no list indent)."""
    p = doc.add_paragraph()
    _apply_spacing(p)
    _tnr_run(p, text or "", bold=False, size_pt=BODY_PT)
    return p


def _label_inline(doc, label, value):
    """Bold label + plain value on same plain paragraph (no indent)."""
    p = doc.add_paragraph()
    _apply_spacing(p)
    _tnr_run(p, f"{label}: ", bold=True, size_pt=BODY_PT)
    _tnr_run(p, value or "", bold=False, size_pt=BODY_PT)
    return p


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_docx(data):
    doc = Document()

    # Page: US Letter, 1" margins (matching template sectPr)
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Set Normal style to Times New Roman 12pt
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(BODY_PT)

    # ── TITLE ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    _apply_spacing(title_para)
    pPr = title_para._p.get_or_add_pPr()
    lvl = OxmlElement("w:outlineLvl")
    lvl.set(qn("w:val"), "0")
    pPr.append(lvl)
    _tnr_run(title_para, "CMO WEEKLY REPORT — MARKETING, PIPELINE & BRAND",
             bold=True, size_pt=TITLE_PT)

    # ── META LINE ──────────────────────────────────────────────────────────────
    meta = doc.add_paragraph()
    _apply_spacing(meta)
    _tnr_run(meta, "Executive: ", bold=True)
    _tnr_run(meta, data.get("executive", "") + "   ")
    _tnr_run(meta, "Week Ending: ", bold=True)
    _tnr_run(meta, data.get("week_ending", "") + "   ")
    _tnr_run(meta, "Reporting To: ", bold=True)
    _tnr_run(meta, data.get("reporting_to", "CEO"))

    # ── SECTION 1: EXECUTIVE SUMMARY ──────────────────────────────────────────
    _section_header(doc, 1, "Marketing Executive Summary")
    _plain_para(doc, "(Topline marketing performance, pipeline impact, brand momentum, strategic insights.)")
    _plain_para(doc, data.get("executive_summary", ""))

    # ── SECTION 2: OBJECTIVES ─────────────────────────────────────────────────
    _section_header(doc, 2, "Marketing Objectives (Planned vs. Completed)")
    for i in range(1, 3):
        p = doc.add_paragraph()
        _apply_spacing(p)
        _tnr_run(p, f"Objective {i}:", bold=True)
        _bullet_item(doc, "Planned", data.get(f"obj{i}_planned", ""))
        _bullet_item(doc, "Completed", data.get(f"obj{i}_completed", ""))
        _bullet_item(doc, "Status", data.get(f"obj{i}_status", ""))

    # ── SECTION 3: KPIs ───────────────────────────────────────────────────────
    _section_header(doc, 3, "Key Marketing Metrics & KPIs")
    kpi_fields = [
        ("pipeline_generated",  "Pipeline generated"),
        ("lead_volume_cac",     "Lead volume / CAC"),
        ("conversion_rates",    "Conversion rates"),
        ("campaign_performance","Campaign performance"),
        ("brand_engagement",    "Brand engagement metrics"),
        ("email_activation",    "Email / list activation"),
        ("content_output",      "Content output"),
    ]
    for key, label in kpi_fields:
        _bullet_item(doc, label, data.get(key, ""))

    # ── SECTION 4: MAJOR WINS ─────────────────────────────────────────────────
    _section_header(doc, 4, "Major Wins & Deliverables")
    wins_text = data.get("major_wins", "")
    lines = [l.strip().lstrip("•-*").strip() for l in wins_text.split("\n") if l.strip()]
    if lines:
        for line in lines:
            _bullet_item(doc, "", line)
    else:
        # two blank bullet placeholders like the template
        _bullet_item(doc, "", "")
        _bullet_item(doc, "", "")

    # ── SECTION 5: STRATEGIC INITIATIVES ──────────────────────────────────────
    _section_header(doc, 5, "Strategic Marketing Initiatives")
    _label_inline(doc, "Initiative", data.get("initiative_name", ""))
    _bullet_item(doc, "Progress This Week", data.get("initiative_progress", ""))
    _bullet_item(doc, "Blockers", data.get("initiative_blockers", ""))
    _bullet_item(doc, "Next Milestone + ETA", data.get("initiative_milestone", ""))

    # ── SECTION 6: CROSS-FUNCTIONAL ───────────────────────────────────────────
    _section_header(doc, 6, "Cross‑Functional Alignment")
    _bullet_item(doc, "Sales enablement delivered", data.get("sales_enablement", ""))
    _bullet_item(doc, "Product marketing updates",  data.get("product_marketing", ""))
    _bullet_item(doc, "Partnerships / PR",           data.get("partnerships_pr", ""))

    # ── SECTION 7: CEO DECISIONS ──────────────────────────────────────────────
    _section_header(doc, 7, "CEO Decisions / Support Needed")
    ceo = data.get("ceo_decisions", "")
    if ceo:
        _bullet_item(doc, "", ceo)
    else:
        _bullet_item(doc, "", "")

    # ── SECTION 8: PRIORITIES NEXT WEEK ───────────────────────────────────────
    _section_header(doc, 8, "Priorities for Next Week")
    prio_text = data.get("priorities_next_week", "")
    prio_lines = [l.strip().lstrip("•-*0123456789.").strip()
                  for l in prio_text.split("\n") if l.strip()]
    if prio_lines:
        for line in prio_lines:
            _bullet_item(doc, "", line)
    else:
        _bullet_item(doc, "", "")

    # ── SECTION 9: CMO COMMENTS ────────────────────────────────────────────────
    _section_header(doc, 9, "CMO Comments & Strategic Notes")
    cmo_comment = data.get("cmo_comment", "").strip()
    cmo_table = doc.add_table(rows=1, cols=1)
    cmo_table.style = "Table Grid"
    cmo_cell = cmo_table.rows[0].cells[0]
    set_cell_bg(cmo_cell, "FFF8E7")
    cmo_cell.paragraphs[0].clear()
    p = cmo_cell.paragraphs[0]
    _apply_spacing(p)
    _tnr_run(p, "CMO Note: ", bold=True, size_pt=BODY_PT)
    _tnr_run(p, cmo_comment if cmo_comment else "[No comment this week]",
             italic=True, size_pt=BODY_PT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/fill-report", methods=["POST"])
def fill_report():
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY not configured on the server"}), 500

    payload = request.json or {}
    dump = (payload.get("dump") or "").strip()
    if not dump:
        return jsonify({"error": "No notes provided"}), 400

    executive = payload.get("executive", "")
    week_ending = payload.get("week_ending", "")
    reporting_to = payload.get("reporting_to", "CEO")

    system = """You are a seasoned Chief Marketing Officer writing your own weekly report.
You've been given raw notes from your week. Turn them into a finished, polished report — the kind a CMO sends to the CEO. No fluff, no filler, no AI-sounding phrases.

Writing rules (non-negotiable):
- Sound like a real executive wrote this, not a chatbot. Confident, direct, specific.
- Every sentence must earn its place. If it doesn't add information, cut it.
- Use active voice. Never say "it was achieved" — say "we hit" or "we closed".
- No corporate fluff: never write "leverage", "synergy", "actionable insights", "robust", "seamlessly", "it is worth noting".
- Numbers and specifics make it credible — keep every number from the notes.
- executive_summary: 2-3 tight sentences. Lead with the biggest result, then momentum, then forward signal. No throat-clearing.
- KPI fields: one crisp line each, e.g. "$1.2M generated, +22% WoW" — not a full sentence.
- obj status fields: plain verdict, e.g. "On Track — 84% of goal" or "At Risk — delayed by agency review".
- major_wins: one punchy line per win, each on its own line. No bullets, no dashes.
- priorities_next_week: one clear action per line. Start each with a verb.
- cmo_comment: the CMO's honest, unfiltered strategic take. Can flag a concern, call out a risk, or note an opportunity. Should sound like a real personal note, not a summary.
- If something isn't mentioned in the notes, leave that field as "".

Return ONLY valid JSON with exactly these keys, nothing else:
{
  "executive_summary": "",
  "obj1_planned": "",
  "obj1_completed": "",
  "obj1_status": "",
  "obj2_planned": "",
  "obj2_completed": "",
  "obj2_status": "",
  "pipeline_generated": "",
  "lead_volume_cac": "",
  "conversion_rates": "",
  "campaign_performance": "",
  "brand_engagement": "",
  "email_activation": "",
  "content_output": "",
  "major_wins": "",
  "initiative_name": "",
  "initiative_progress": "",
  "initiative_blockers": "",
  "initiative_milestone": "",
  "sales_enablement": "",
  "product_marketing": "",
  "partnerships_pr": "",
  "ceo_decisions": "",
  "priorities_next_week": "",
  "cmo_comment": ""
}"""

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        system=system,
        messages=[{"role": "user", "content": f"Weekly notes:\n{dump}"}],
    )

    raw = message.content[0].text.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        fields = json.loads(raw)
    except Exception:
        return jsonify({"error": "AI returned unexpected format. Try again."}), 500

    # Merge in the header fields
    fields["executive"] = executive
    fields["week_ending"] = week_ending
    fields["reporting_to"] = reporting_to or "CEO"

    return jsonify(fields)


@app.route("/api/download", methods=["POST"])
def download():
    data = request.json or {}
    buf = generate_docx(data)
    week = data.get("week_ending", "report").replace("/", "-").replace(" ", "_")
    filename = f"CMO_Weekly_Report_{week}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    print("The Awareness Group — CMO Report Generator running at http://localhost:5151")
    app.run(debug=False, port=5151)
